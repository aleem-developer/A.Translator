import os
import math
import itertools
import pdfplumber
import warnings
import threading
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt
from pathlib import Path
from PIL import Image, ImageTk
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as tb
from ttkbootstrap.constants import *
import sys

# --- Handle paths for PyInstaller ---
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS  # PyInstaller temp folder
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

warnings.filterwarnings("ignore", category=UserWarning, module="pdfplumber")

LANGUAGES = {
    "en": "English", "es": "Spanish", "fr": "French", "de": "German",
    "it": "Italian", "pt": "Portuguese", "ru": "Russian", "ar": "Arabic",
    "hi": "Hindi", "bn": "Bengali", "zh-CN": "Chinese (Simplified)", "ja": "Japanese",
    "ko": "Korean", "tr": "Turkish", "nl": "Dutch", "pl": "Polish",
    "vi": "Vietnamese", "fa": "Persian", "sv": "Swedish", "uk": "Ukrainian",
    "id": "Indonesian", "cs": "Czech", "ro": "Romanian", "el": "Greek",
    "th": "Thai", "ta": "Tamil", "te": "Telugu", "ms": "Malay", "ur": "Urdu",
    "he": "Hebrew"
}

class TranslatorWaveLoader(tk.Canvas):
    def __init__(self, parent, width=420, height=30, fg='#00BFFF', bg='#1a1a1a', **kwargs):
        super().__init__(parent, width=width, height=height, bg=bg, highlightthickness=0, **kwargs)
        self.width = width
        self.height = height
        self.fg = fg
        self.progress = 0
        self.offset = 0
        self.translating_texts = itertools.cycle([
            "Translating...", "Traduciendo...", "Traduction en cours...",
            "Übersetzen...", "ترجمہ ہو رہا ہے...", "翻译中...", "正在翻译...",
            "ترجمه در حال انجام...", "ترجمہ کر رہا ہے..."
        ])
        self.translating_label = tk.Label(parent, text="", font=("Segoe UI", 12, "bold"), fg='cyan', bg=bg)
        self.translating_label.pack(pady=(8, 4))
        self.percent_label = tk.Label(parent, text="0%", font=("Segoe UI", 10, "bold"), fg='white', bg=bg)
        self.percent_label.pack(pady=(0, 10))
        self._animate()
        self._rotate_text()

    def set_progress(self, percent):
        self.progress = max(0, min(100, percent))
        self.percent_label.config(text=f"{int(self.progress)}%")

    def _animate(self):
        self.delete("wave")
        fill_height = self.height * (self.progress / 100)
        for i, color in enumerate(["#00BFFF", "#1E90FF", "#4169E1"]):
            points = []
            for x in range(0, self.width + 2, 2):
                angle = (x + self.offset + i * 15) * 0.05
                y = self.height - fill_height + 5 * math.sin(angle)
                points.append((x, y))
            for x, y in points:
                self.create_line(x, self.height, x, y, fill=color, tags="wave")
        self.offset += 2
        self.after(50, self._animate)

    def _rotate_text(self):
        next_text = next(self.translating_texts)
        self.translating_label.config(text=next_text)
        self.after(1000, self._rotate_text)

def show_splash():
    splash_root = tk.Tk()
    splash_root.overrideredirect(True)
    splash_root.configure(bg='black')
    width, height = 500, 300
    screen_width = splash_root.winfo_screenwidth()
    screen_height = splash_root.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    splash_root.geometry(f"{width}x{height}+{x}+{y}")
    img = Image.open(resource_path("a.logo.png"))
    img = img.resize((width, height), Image.LANCZOS)
    splash_img = ImageTk.PhotoImage(img)
    label = tk.Label(splash_root, image=splash_img, border=0)
    label.image = splash_img
    label.pack()
    splash_root.after(2500, splash_root.destroy)
    splash_root.mainloop()

class PDFTranslatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF Converter - by Aleem.Developer")
        self.root.geometry("860x660+{}+{}".format((root.winfo_screenwidth() - 860) // 2, (root.winfo_screenheight() - 660) // 2))
        self.root.resizable(False, False)
        self.style = tb.Style("darkly")
        self.mode = tk.StringVar(value="1")
        self.lang_code = tk.StringVar()
        self.file_path = ""
        self.output_dir = tk.StringVar(value=str(Path.home() / "Documents"))
        self.build_ui()

    def build_ui(self):
        tb.Label(self.root, text="📘 PDF Language Converter", font=("Segoe UI", 24, "bold"), bootstyle="info inverse", foreground="white").pack(pady=20)
        tb.Label(self.root, text="by Aleem.Developer", font=("Segoe UI", 12), bootstyle="secondary", foreground="white").pack()
        file_frame = tb.Frame(self.root)
        file_frame.pack(pady=10)
        tb.Button(file_frame, text="📂 Select PDF File", bootstyle=PRIMARY, command=self.select_file).grid(row=0, column=0, padx=10)
        self.file_label = tb.Label(file_frame, text="No file selected", font=("Segoe UI", 11), foreground="white")
        self.file_label.grid(row=0, column=1)
        settings = tb.Labelframe(self.root, text="Conversion Settings", bootstyle=INFO, padding=20)
        settings.pack(pady=20, padx=20, fill=X)
        tb.Label(settings, text="Output Mode:", font=("Segoe UI", 12), foreground="white").grid(row=0, column=0, sticky="w")
        tb.Radiobutton(settings, text="English Only", variable=self.mode, value="1", command=self.toggle_lang).grid(row=1, column=0, sticky="w")
        tb.Radiobutton(settings, text="Translated Only", variable=self.mode, value="2", command=self.toggle_lang).grid(row=2, column=0, sticky="w")
        tb.Radiobutton(settings, text="English + Translated", variable=self.mode, value="3", command=self.toggle_lang).grid(row=3, column=0, sticky="w")
        tb.Label(settings, text="Target Language:", font=("Segoe UI", 12), foreground="white").grid(row=0, column=1, sticky="w", padx=(50, 0))
        self.lang_combo = tb.Combobox(settings, textvariable=self.lang_code, width=30, state="disabled", values=[f"{code} = {name}" for code, name in LANGUAGES.items()])
        self.lang_combo.grid(row=1, column=1, rowspan=2, padx=(50, 0), pady=5)
        tb.Label(settings, text="Output Folder:", font=("Segoe UI", 12), foreground="white").grid(row=4, column=0, sticky="w", pady=(10, 0))
        out_frame = tb.Frame(settings)
        out_frame.grid(row=5, column=0, columnspan=2, sticky="w", pady=5)
        tb.Entry(out_frame, textvariable=self.output_dir, width=60).grid(row=0, column=0, padx=(0, 10))
        tb.Button(out_frame, text="Change Folder", bootstyle=SECONDARY, command=self.change_folder).grid(row=0, column=1)
        tb.Button(self.root, text="🚀 Start Conversion", bootstyle=SUCCESS, command=self.start_conversion_thread, width=30).pack(pady=20)
        self.loader = TranslatorWaveLoader(self.root)
        self.loader.pack()
        self.open_button = tb.Button(self.root, text="📂 Open Output Folder", bootstyle=INFO, command=self.open_output_folder, state="disabled")
        self.open_button.pack(pady=5)
        tb.Label(self.root, text="© 2025 Aleem.Developer. All rights reserved.", font=("Segoe UI", 10), bootstyle="secondary", foreground="white").pack(pady=(10, 0))

    def toggle_lang(self):
        if self.mode.get() in ["2", "3"]:
            self.lang_combo.config(state="readonly")
        else:
            self.lang_combo.set("")
            self.lang_combo.config(state="disabled")

    def select_file(self):
        self.file_path = filedialog.askopenfilename(title="Select PDF File", filetypes=[("PDF Files", "*.pdf")])
        if self.file_path:
            self.file_label.config(text=f"📌 {os.path.basename(self.file_path)}", bootstyle="success")
        else:
            self.file_label.config(text="❌ No file selected", bootstyle="danger")

    def change_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.output_dir.set(folder)

    def start_conversion_thread(self):
        self.loader.set_progress(0)
        self.open_button.config(state="disabled")
        threading.Thread(target=self.convert_pdf_text).start()

    def open_output_folder(self):
        os.startfile(self.output_dir.get())

    def convert_pdf_text(self):
        if not self.file_path:
            messagebox.showerror("Error", "Please select a PDF file.")
            return
        mode = self.mode.get()
        lang_code = None
        lang_name = ""
        if mode in ['2', '3']:
            if not self.lang_code.get():
                messagebox.showerror("Error", "Please select a language.")
                return
            lang_code = self.lang_code.get().split(" = ")[0]
            lang_name = LANGUAGES[lang_code]
        filename = "Extracted_English.docx" if mode == '1' else f"Translated_{lang_name}.docx" if mode == '2' else f"Dual_Language_{lang_name}.docx"
        output_path = os.path.join(self.output_dir.get(), filename)
        doc = Document()
        try:
            with pdfplumber.open(self.file_path) as pdf:
                total_pages = len(pdf.pages)
                for idx, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if not text:
                        continue
                    if mode == '1':
                        para = doc.add_paragraph(text.strip())
                        para.alignment = 0
                        para.runs[0].font.size = Pt(11)
                    elif mode == '2':
                        translated = GoogleTranslator(source='auto', target=lang_code).translate(text)
                        para = doc.add_paragraph(translated.strip())
                        para.alignment = 2 if lang_code in ["ur", "ar", "he"] else 0
                        para.runs[0].font.size = Pt(12)
                    elif mode == '3':
                        translated = GoogleTranslator(source='auto', target=lang_code).translate(text)
                        en = doc.add_paragraph("English:\n" + text.strip())
                        en.alignment = 0
                        en.runs[0].font.size = Pt(11)
                        doc.add_paragraph("")
                        tr = doc.add_paragraph(f"{lang_name}:\n" + translated.strip())
                        tr.alignment = 2 if lang_code in ["ur", "ar", "he"] else 0
                        tr.runs[0].font.size = Pt(12)
                    self.loader.set_progress((idx + 1) / total_pages * 100)
                    if idx + 1 < total_pages:
                        doc.add_page_break()
            doc.save(output_path)
            messagebox.showinfo("Success", f"✅ Output saved to:\n{output_path}\n\nCreated by Aleem.Developer")
            self.open_button.config(state="normal")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred:\n{str(e)}")

if __name__ == "__main__":
    show_splash()
    root = tb.Window(themename="darkly")
    root.iconbitmap(resource_path("a.logo.ico"))
    app = PDFTranslatorApp(root)
    root.mainloop()
